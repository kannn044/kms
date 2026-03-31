import streamlit as st
import pandas as pd
import time
import os
from modules.knowledge import (
    add_knowledge_item,
    update_knowledge_item,
    delete_knowledge_item,
    get_knowledge_item,
    search_knowledge_items,
    get_categories
)
from modules.vectorstore import add_to_vectorstore
from modules.ui.styles import (
    inject_global_css, render_page_header, render_knowledge_card,
    render_empty_state, render_badge
)


def show_add_knowledge_page(conn):
    """Display form to add new knowledge with OCR capability."""
    inject_global_css()

    if conn is None:
        from modules.database import get_db_connection
        conn = get_db_connection()

    render_page_header(
        title="Add New Knowledge",
        subtitle="Share your expertise with the organization",
        icon="➕"
    )

    # Tabs
    add_tabs = st.tabs(["✏️ Manual Entry", "📄 Extract from File (OCR)"])

    with add_tabs[0]:
        show_manual_entry_form(conn)

    with add_tabs[1]:
        show_ocr_extraction_form(conn)


def show_manual_entry_form(conn):
    """Manual knowledge entry form with improved layout."""
    with st.form("manual_knowledge_form"):
        title = st.text_input("📌 Title *", placeholder="Enter a descriptive title")
        content = st.text_area("📝 Content *", height=200, placeholder="Write your knowledge content here...")

        col1, col2 = st.columns(2)
        with col1:
            existing_categories = get_categories()
            default_categories = ["General", "Technical", "Process", "Project", "Research"]
            category_options = sorted(set(default_categories + existing_categories))
            category = st.selectbox("📂 Category *", category_options)
        with col2:
            tags = st.text_input("🏷️ Tags", placeholder="tag1, tag2, tag3")

        uploaded_file = st.file_uploader(
            "📎 Attach File (optional)",
            type=["pdf", "docx", "xlsx", "txt", "jpg", "png", "jpeg"]
        )

        submit_button = st.form_submit_button("💾 Save Knowledge", use_container_width=True)

    if submit_button:
        if not title or not content or not category:
            st.error("❌ Title, Content and Category are required")
            return

        with st.spinner("Saving and indexing knowledge..."):
            progress = st.progress(0)
            status_text = st.empty()

            status_text.text("📝 Saving content...")
            progress.progress(30)
            time.sleep(0.3)

            status_text.text("🧠 Creating vector embeddings...")
            progress.progress(60)
            time.sleep(0.3)

            status_text.text("📊 Updating vector database...")
            progress.progress(90)
            time.sleep(0.3)

            author_id = st.session_state['user_id']
            success, result = add_knowledge_item(title, content, category, tags, author_id, uploaded_file)

            if success:
                progress.progress(100)
                status_text.success("✅ Knowledge Added Successfully!")
                time.sleep(1)
                st.rerun()
            else:
                progress.empty()
                st.error(f"❌ Failed to add knowledge: {result}")


def show_ocr_extraction_form(conn):
    """Extract text from uploaded files using OCR."""
    import io

    _missing = []
    try:
        from PIL import Image
    except ImportError:
        Image = None
        _missing.append("`pillow`")

    try:
        import fitz
    except ImportError:
        fitz = None
        _missing.append("`pymupdf`")

    try:
        import docx
    except ImportError:
        docx = None
        _missing.append("`python-docx`")

    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        _tesseract_ok = True
    except ImportError:
        pytesseract = None
        _tesseract_ok = False
        _missing.append("`pytesseract`")
    except Exception:
        try:
            import pytesseract as pytesseract
        except ImportError:
            pytesseract = None
        _tesseract_ok = False

    if _missing:
        st.error(
            "Some required libraries are missing. Install them with:\n\n"
            "```\npip install pytesseract pillow pymupdf python-docx\n```\n\n"
            f"Missing: {', '.join(_missing)}"
        )
        if len(_missing) == 4:
            return

    _supported_types = []
    if fitz is not None:
        _supported_types.append("pdf")
    if docx is not None:
        _supported_types += ["docx"]
    if Image is not None and pytesseract is not None:
        _supported_types += ["jpg", "jpeg", "png"]

    if not _supported_types:
        st.error("No file types can be processed. Please install the required libraries above.")
        return

    if not _tesseract_ok:
        st.warning(
            "**Tesseract OCR engine not found.** Image OCR (JPG/PNG) is disabled.\n\n"
            "Install Tesseract on macOS: `brew install tesseract`\n\n"
            "PDF and DOCX text extraction still work without Tesseract."
        )

    st.info("📄 Upload an image or document to extract text automatically.")

    uploaded_ocr_file = st.file_uploader(
        "Upload file for text extraction",
        type=_supported_types,
        key="ocr_file_uploader"
    )

    extracted_text = ""
    extracted_title = ""
    ocr_status = st.empty()
    ocr_progress = st.empty()

    if uploaded_ocr_file:
        with st.spinner("Processing file..."):
            progress_bar = ocr_progress.progress(0)
            ocr_status.info("🔍 Analyzing file type...")
            time.sleep(0.3)

            file_extension = uploaded_ocr_file.name.split('.')[-1].lower()

            try:
                if file_extension in ['jpg', 'jpeg', 'png']:
                    if not _tesseract_ok:
                        ocr_status.error("Tesseract OCR is not installed.")
                        progress_bar.empty()
                        return
                    ocr_status.info("🔍 Performing OCR on image...")
                    progress_bar.progress(30)
                    image = Image.open(uploaded_ocr_file)
                    try:
                        extracted_text = pytesseract.image_to_string(image, lang='tha+eng')
                    except Exception:
                        extracted_text = pytesseract.image_to_string(image)
                    progress_bar.progress(80)

                elif file_extension == 'pdf':
                    ocr_status.info("📕 Extracting text from PDF...")
                    progress_bar.progress(20)
                    pdf_bytes = uploaded_ocr_file.read()
                    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
                    total_pages = len(pdf_document)
                    all_text = []
                    for i, page in enumerate(pdf_document):
                        text = page.get_text()
                        all_text.append(text)
                        progress_bar.progress(20 + int(60 * (i + 1) / total_pages))
                        ocr_status.info(f"📕 Processing page {i+1} of {total_pages}...")
                    extracted_text = "\n\n".join(all_text)

                elif file_extension == 'docx':
                    ocr_status.info("📘 Extracting from Word document...")
                    progress_bar.progress(30)
                    doc = docx.Document(io.BytesIO(uploaded_ocr_file.read()))
                    all_paragraphs = []
                    total_paras = len(doc.paragraphs)
                    for i, para in enumerate(doc.paragraphs):
                        if para.text:
                            all_paragraphs.append(para.text)
                        if i % 10 == 0:
                            progress_bar.progress(30 + int(50 * (i + 1) / total_paras))
                    extracted_text = "\n\n".join(all_paragraphs)

                progress_bar.progress(90)

                if extracted_text:
                    lines = extracted_text.split('\n')
                    extracted_title = next((line for line in lines if line.strip()), "")
                    if len(extracted_title) > 70:
                        extracted_title = extracted_title[:67] + "..."
                    progress_bar.progress(100)
                    ocr_status.success("✅ Text successfully extracted!")
                else:
                    ocr_status.warning("⚠️ No text could be extracted from this file.")

            except Exception as e:
                ocr_status.error(f"❌ Error extracting text: {e}")
                progress_bar.empty()

    if extracted_text:
        with st.form("ocr_knowledge_form"):
            st.markdown("### Review Extracted Content")

            title = st.text_input("📌 Title *", value=extracted_title)
            content = st.text_area("📝 Content *", value=extracted_text, height=300)

            col1, col2 = st.columns(2)
            with col1:
                existing_categories = get_categories()
                default_categories = ["General", "Technical", "Process", "Project", "Research"]
                category_options = sorted(set(default_categories + existing_categories))
                category = st.selectbox("📂 Category *", category_options, key="ocr_category")
            with col2:
                tags = st.text_input("🏷️ Tags", key="ocr_tags", placeholder="tag1, tag2, tag3")

            keep_original = st.checkbox("📎 Attach original file", value=True)
            submit_ocr = st.form_submit_button("💾 Save Knowledge", use_container_width=True)

        if submit_ocr:
            if not title or not content or not category:
                st.error("❌ Title, Content and Category are required")
                return

            with st.spinner("Saving extracted knowledge..."):
                save_progress = st.progress(0)
                save_status = st.empty()

                save_status.text("📝 Processing content...")
                save_progress.progress(30)
                time.sleep(0.3)

                save_status.text("🧠 Creating vector embeddings...")
                save_progress.progress(60)
                time.sleep(0.3)

                file_to_attach = uploaded_ocr_file if keep_original else None
                author_id = st.session_state['user_id']
                success, result = add_knowledge_item(title, content, category, tags, author_id, file_to_attach)

                if success:
                    save_progress.progress(100)
                    save_status.success("✅ Knowledge Added Successfully!")
                    time.sleep(1)
                    st.rerun()
                else:
                    save_progress.empty()
                    st.error(f"❌ Failed to add knowledge: {result}")


def show_manage_knowledge_page():
    """Display styled knowledge management interface."""
    inject_global_css()

    render_page_header(
        title="Manage Knowledge",
        subtitle="Edit, update, or delete existing knowledge items",
        icon="⚙️"
    )

    all_items = search_knowledge_items()

    if all_items:
        # Show items as styled cards with select
        st.markdown(f"""
        <div style="color: #8899AA; margin-bottom: 1rem;">
            📊 Total: <strong style="color: #E8ECF1;">{len(all_items)}</strong> knowledge items
        </div>
        """, unsafe_allow_html=True)

        for item in all_items:
            author_name = item.get('full_name') or item.get('username', 'Unknown')
            st.markdown(
                render_knowledge_card(
                    title=f"[ID: {item['id']}] {item['title']}",
                    category=item['category'],
                    tags=item.get('tags', ''),
                    date=item['last_updated'],
                    author=author_name
                ),
                unsafe_allow_html=True
            )

        st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

        # Select item to edit
        item_ids = [item['id'] for item in all_items]
        item_labels = {item['id']: f"[{item['id']}] {item['title']}" for item in all_items}

        selected_id = st.selectbox(
            "Select item to edit/delete",
            options=item_ids,
            format_func=lambda x: item_labels.get(x, str(x))
        )

        if st.button("📝 Load Item", use_container_width=True):
            item = get_knowledge_item(selected_id)
            if item:
                st.session_state["edit_id"] = item['id']
                st.session_state["edit_title"] = item['title']
                st.session_state["edit_content"] = item['content']
                st.session_state["edit_category"] = item['category']
                st.session_state["edit_tags"] = item['tags']
                st.session_state["edit_file_path"] = item['file_path']
                st.rerun()
            else:
                st.error("❌ Item not found")

        if "edit_id" in st.session_state:
            show_edit_form()
    else:
        st.markdown(
            render_empty_state("📭", "No knowledge items found", "Start adding some knowledge!"),
            unsafe_allow_html=True
        )


def show_edit_form():
    """Edit form for a loaded knowledge item."""
    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
    st.markdown("### ✏️ Edit Knowledge Item")

    with st.form("edit_knowledge_form"):
        edit_title = st.text_input("📌 Title", value=st.session_state["edit_title"])
        edit_content = st.text_area("📝 Content", value=st.session_state["edit_content"], height=200)

        col1, col2 = st.columns(2)
        with col1:
            existing_categories = get_categories()
            default_categories = ["General", "Technical", "Process", "Project", "Research"]
            category_options = sorted(set(default_categories + existing_categories))
            try:
                category_index = category_options.index(st.session_state["edit_category"])
            except ValueError:
                category_index = 0
            edit_category = st.selectbox("📂 Category", category_options, index=category_index)

        with col2:
            edit_tags = st.text_input("🏷️ Tags", value=st.session_state["edit_tags"] or "")

        if st.session_state["edit_file_path"]:
            file_name = os.path.basename(st.session_state["edit_file_path"])
            st.info(f"📎 Current file: {file_name}")

        new_file = st.file_uploader(
            "📎 Replace File (optional)",
            type=["pdf", "docx", "xlsx", "txt", "jpg", "png"]
        )

        col1, col2 = st.columns(2)
        with col1:
            update_button = st.form_submit_button("💾 Update", use_container_width=True)
        with col2:
            delete_button = st.form_submit_button("🗑️ Delete", use_container_width=True)

    if update_button:
        if not edit_title or not edit_content or not edit_category:
            st.error("❌ Title, Content and Category are required")
            return

        with st.spinner("Updating knowledge..."):
            progress = st.progress(0)
            status = st.empty()

            status.text("📝 Updating content...")
            progress.progress(30)
            time.sleep(0.3)

            status.text("🧠 Updating embeddings...")
            progress.progress(60)
            time.sleep(0.3)

            success, message = update_knowledge_item(
                st.session_state["edit_id"],
                edit_title, edit_content, edit_category, edit_tags, new_file
            )

            if success:
                progress.progress(100)
                status.success("✅ Knowledge Updated Successfully!")
                for key in list(st.session_state.keys()):
                    if key.startswith("edit_"):
                        del st.session_state[key]
                time.sleep(1)
                st.rerun()
            else:
                st.error(f"❌ Failed to update: {message}")

    if delete_button:
        with st.spinner("Deleting knowledge..."):
            progress = st.progress(0)
            status = st.empty()

            status.text("🗑️ Removing entry...")
            progress.progress(40)
            time.sleep(0.3)

            status.text("🔄 Rebuilding vector index...")
            progress.progress(70)
            time.sleep(0.3)

            success, message = delete_knowledge_item(st.session_state["edit_id"])

            if success:
                progress.progress(100)
                status.success("✅ Knowledge Deleted Successfully!")
                for key in list(st.session_state.keys()):
                    if key.startswith("edit_"):
                        del st.session_state[key]
                time.sleep(1)
                st.rerun()
            else:
                st.error(f"❌ Failed to delete: {message}")